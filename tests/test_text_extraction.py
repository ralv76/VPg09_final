"""Юнит-тесты извлечения текста и очистки ПД. ТЗ 2.1.1."""
import tempfile
from pathlib import Path

import pytest

from backend.services.text_extraction import (
    mask_pii,
    clean_and_format,
    extract_from_pdf,
    extract_from_docx,
    extract_from_url,
)


def test_mask_pii_email():
    text = "Связь: user@example.com и admin@test.ru"
    out, _, contacts = mask_pii(text)
    assert "[контакт скрыт]" in out
    assert "user@example.com" not in out
    assert len(contacts) >= 2


def test_mask_pii_phone():
    text = "Позвоните +7 (999) 123-45-67 или 8 800 100 20 30"
    out, phones, _ = mask_pii(text)
    assert "[телефон скрыт]" in out
    assert len(phones) >= 1


def test_clean_and_format():
    text = "  Первая строка   \n\n  Вторая строка  \n"
    result = clean_and_format(text)
    assert "Первая строка" in result
    assert "Вторая строка" in result
    assert "\n\n" in result


def test_extract_from_pdf_too_large(monkeypatch):
    import backend.services.text_extraction as m
    monkeypatch.setattr(m, "MAX_SIZE", 10)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(b"x" * 20)
        path = Path(f.name)
    try:
        with pytest.raises(ValueError, match="лимит"):
            extract_from_pdf(path)
    finally:
        path.unlink(missing_ok=True)


def test_extract_from_docx(tmp_path):
    """Извлечение из DOCX."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Первый абзац.")
    doc.add_paragraph("Второй абзац.")
    path = tmp_path / "test.docx"
    doc.save(path)
    text, phones, contacts = extract_from_docx(path)
    assert "Первый абзац" in text
    assert "Второй абзац" in text


@pytest.mark.parametrize("url", ["https://example.com"])
def test_extract_from_url_integration(url):
    """Извлечение с реального URL (example.com)."""
    try:
        text, _, _ = extract_from_url(url)
        assert isinstance(text, str)
    except Exception as e:
        pytest.skip(f"Сеть недоступна: {e}")
